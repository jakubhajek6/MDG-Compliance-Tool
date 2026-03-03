"""
Modul 3 – AML kontroly
Automatická AML prověrka fyzické nebo právnické osoby.
"""

import json
import time
from datetime import datetime
from io import BytesIO

import streamlit as st
from docx import Document

from modules.aml_checks import run_aml_check
from db.database import init_db, save_aml_check, get_aml_checks, log_audit

# ===== PAGE CONFIG =====
st.set_page_config(page_title="MDG – AML kontroly", page_icon="🔍", layout="wide")

init_db()

# ===== CSS =====
PRIMARY = "#1B3A6B"
CSS = f"""
<style>
.stButton > button, .stDownloadButton > button {{
  background-color: {PRIMARY} !important; color: white !important; border: 1px solid {PRIMARY} !important;
}}
div.stProgress > div > div {{ background-color: {PRIMARY} !important; }}
.breadcrumb {{ color: #888; font-size: 0.85rem; margin-bottom: 0.5rem; }}
.aml-clean {{ background-color: #d4edda; border-radius: 8px; padding: 12px; margin: 8px 0; }}
.aml-warning {{ background-color: #fff3cd; border-radius: 8px; padding: 12px; margin: 8px 0; }}
.aml-hit {{ background-color: #f8d7da; border-radius: 8px; padding: 12px; margin: 8px 0; }}
</style>
"""
st.markdown(CSS, unsafe_allow_html=True)

# ===== HEADER =====
st.markdown('<div class="breadcrumb">Domů / AML kontroly</div>', unsafe_allow_html=True)
st.markdown("## 🔍 AML kontroly")
st.markdown("Automatická AML prověrka fyzické nebo právnické osoby proti sankčním seznamům, PEP databázi a insolvenčnímu rejstříku.")

# ===== Předvyplnění z jiných modulů =====
prefill_name = st.session_state.pop("aml_prefill_name", "")
prefill_type = st.session_state.pop("aml_prefill_type", "FO")

st.markdown("---")

# ===== Vstupy =====
col1, col2, col3 = st.columns([3, 2, 2])
with col1:
    entity_name = st.text_input("Jméno osoby / název firmy", value=prefill_name,
                                placeholder="např. Jan Novák nebo ABC s.r.o.")
with col2:
    ico = st.text_input("IČO (volitelné, pro PO)", placeholder="12345678")
with col3:
    entity_type = st.selectbox("Typ entity", ["FO (fyzická osoba)", "PO (právnická osoba)"],
                               index=0 if prefill_type == "FO" else 1)

entity_type_code = "FO" if entity_type.startswith("FO") else "PO"

# ===== Spuštění kontroly =====
if st.button("🔍 Spustit AML kontrolu", type="primary"):
    if not entity_name.strip():
        st.error("Zadejte jméno osoby nebo název firmy.")
    else:
        st.markdown("---")
        st.subheader("Průběh kontroly")

        # Progress container
        progress_bar = st.progress(0)
        status_container = st.container()
        step_statuses = []

        def progress_callback(step_name, status, detail, progress):
            progress_bar.progress(max(0, min(100, int(progress * 100))))
            step_statuses.append({"step": step_name, "status": status, "detail": detail})

        # Spustit kontrolu
        with st.spinner("Probíhá AML kontrola..."):
            results = run_aml_check(
                name=entity_name.strip(),
                ico=ico.strip(),
                entity_type=entity_type_code,
                progress_cb=progress_callback,
            )

        progress_bar.progress(100)

        # Uložit do DB
        save_aml_check(
            ico=ico.strip(),
            entity_name=entity_name.strip(),
            entity_type=entity_type_code,
            result_status=results["overall_status"],
            details=results,
            risk_score=0,
        )
        log_audit("AML", "check", ico=ico.strip(), entity_name=entity_name.strip(),
                  details=f"status={results['overall_status']}")

        # ===== Výsledky =====
        st.markdown("---")
        st.subheader("Výsledky AML kontroly")

        # Celkový souhrn
        overall = results["overall_status"]
        if overall == "clean":
            st.markdown('<div class="aml-clean">✅ <strong>Výsledek: ČISTÝ</strong> – Žádné nálezy v kontrolovaných zdrojích.</div>', unsafe_allow_html=True)
        elif overall == "warning":
            st.markdown('<div class="aml-warning">⚠️ <strong>Výsledek: K OVĚŘENÍ</strong> – Nalezeny záznamy vyžadující manuální kontrolu.</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="aml-hit">🚨 <strong>Výsledek: HIT</strong> – Nalezeny záznamy na sankčních seznamech!</div>', unsafe_allow_html=True)

        st.markdown(f"**Kontrolovaná osoba:** {results['entity_name']}")
        st.markdown(f"**Typ:** {results['entity_type']}")
        if results.get("ico"):
            st.markdown(f"**IČO:** {results['ico']}")
        st.markdown(f"**Datum kontroly:** {results['check_date']}")
        st.markdown(f"**Celkový počet nálezů:** {results['total_hits']}")

        # Detailní výsledky po krocích
        st.markdown("---")
        st.subheader("Detailní výsledky")

        for check in results.get("checks", []):
            status = check.get("status", "clean")
            hits = check.get("hits", 0)
            name = check.get("name", "")

            if status == "clean":
                icon = "✅"
                color = "#28a745"
            elif status == "warning":
                icon = "⚠️"
                color = "#ffc107"
            elif status == "hit":
                icon = "🚨"
                color = "#dc3545"
            else:
                icon = "❓"
                color = "#6c757d"

            st.markdown(f"### {icon} {name}")
            st.markdown(f"**Status:** :{color}[{status.upper()}] | **Počet nálezů:** {hits}")

            details = check.get("details", [])
            if details:
                for d in details[:5]:
                    if d.get("error"):
                        st.error(f"Chyba: {d['error']}")
                    elif d.get("matched_name"):
                        sim = d.get("similarity", 0)
                        st.markdown(f"- **{d['matched_name']}** (shoda: {sim}%)"
                                    + (f" – role: {d.get('role', '')}" if d.get("role") else "")
                                    + (f" – zdroj: {d.get('source', '')}" if d.get("source") else ""))
                    elif d.get("text"):
                        st.markdown(f"- {d['text'][:200]}")
                    elif d.get("note"):
                        st.info(d["note"])
            elif status == "clean":
                st.markdown("Žádné nálezy.")

            st.markdown("")

        # ===== Export AML reportu =====
        st.markdown("---")
        st.subheader("Export")

        def generate_aml_report_docx(results):
            doc = Document()
            doc.add_heading("AML REPORT", level=1)
            doc.add_paragraph(f"Datum: {results['check_date']}")
            doc.add_paragraph(f"Kontrolovaná osoba: {results['entity_name']}")
            doc.add_paragraph(f"Typ: {results['entity_type']}")
            if results.get("ico"):
                doc.add_paragraph(f"IČO: {results['ico']}")
            doc.add_paragraph("")

            overall = results["overall_status"]
            if overall == "clean":
                doc.add_paragraph("CELKOVÝ VÝSLEDEK: ČISTÝ")
            elif overall == "warning":
                doc.add_paragraph("CELKOVÝ VÝSLEDEK: K OVĚŘENÍ")
            else:
                doc.add_paragraph("CELKOVÝ VÝSLEDEK: HIT")

            doc.add_paragraph(f"Celkový počet nálezů: {results['total_hits']}")
            doc.add_paragraph("")

            for check in results.get("checks", []):
                doc.add_heading(check.get("name", ""), level=2)
                doc.add_paragraph(f"Status: {check.get('status', '').upper()}")
                doc.add_paragraph(f"Počet nálezů: {check.get('hits', 0)}")
                for d in check.get("details", [])[:10]:
                    if d.get("matched_name"):
                        doc.add_paragraph(f"  - {d['matched_name']} (shoda: {d.get('similarity', 0)}%)", style="List Bullet")
                    elif d.get("text"):
                        doc.add_paragraph(f"  - {d['text'][:200]}", style="List Bullet")
                doc.add_paragraph("")

            doc.add_paragraph("")
            doc.add_paragraph("Generováno MDG Compliance Tool")

            buf = BytesIO()
            doc.save(buf)
            return buf.getvalue()

        report_bytes = generate_aml_report_docx(results)
        st.download_button(
            "📄 Stáhnout AML report (.docx)",
            data=report_bytes,
            file_name=f"aml_report_{entity_name.strip().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
        )

        # JSON export
        st.download_button(
            "📋 Stáhnout detaily (JSON)",
            data=json.dumps(results, ensure_ascii=False, indent=2),
            file_name=f"aml_details_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            mime="application/json",
        )

        # Uložit do session pro Modul 7
        st.session_state["last_aml_results"] = results

# ===== Historie kontrol =====
st.markdown("---")
st.subheader("Historie AML kontrol")

history = get_aml_checks(limit=20)
if history:
    for h in history:
        status = h.get("result_status", "")
        icon = "✅" if status == "clean" else ("⚠️" if status == "warning" else "🚨")
        st.markdown(
            f"{icon} **{h.get('entity_name', 'N/A')}** "
            f"({h.get('entity_type', '')}) – "
            f"IČO: {h.get('ico', 'N/A')} – "
            f"{h.get('check_date', 'N/A')} – "
            f"Status: {status.upper()}"
        )
else:
    st.info("Zatím nebyly provedeny žádné AML kontroly.")
