"""
Generování Word dokumentů z šablon pomocí python-docx.
Nahrazuje {{PLACEHOLDER}} tagy reálnými daty.
"""

import os
import re
import zipfile
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt

TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "data" / "templates"


def get_available_templates() -> list[dict]:
    """Vrátí seznam dostupných šablon."""
    templates = []
    if not TEMPLATES_DIR.exists():
        return templates

    for f in sorted(TEMPLATES_DIR.glob("*.docx")):
        templates.append({
            "name": f.stem,
            "filename": f.name,
            "path": str(f),
        })
    return templates


def build_placeholders(company_info: dict) -> dict:
    """Sestaví slovník placeholderů z dat firmy."""
    now = datetime.now()

    # Jednatel – první statutární orgán typu FO
    jednatel_jmeno = ""
    jednatel_funkce = ""
    for stat in (company_info.get("statutarni_organ") or []):
        if stat.get("typ") == "FO":
            jednatel_jmeno = stat.get("jmeno", "")
            jednatel_funkce = stat.get("funkce", "")
            break
    if not jednatel_jmeno:
        for stat in (company_info.get("statutarni_organ") or []):
            jednatel_jmeno = stat.get("jmeno", "")
            jednatel_funkce = stat.get("funkce", "")
            break

    return {
        "{{NAZEV_FIRMY}}": company_info.get("nazev", ""),
        "{{ICO}}": company_info.get("ico", ""),
        "{{DIC}}": company_info.get("dic", ""),
        "{{SIDLO_ULICE}}": company_info.get("sidlo_ulice", ""),
        "{{SIDLO_MESTO}}": company_info.get("sidlo_mesto", ""),
        "{{SIDLO_PSC}}": company_info.get("sidlo_psc", ""),
        "{{PRAVNI_FORMA}}": company_info.get("pravni_forma", ""),
        "{{JEDNATEL_JMENO}}": jednatel_jmeno,
        "{{JEDNATEL_FUNKCE}}": jednatel_funkce,
        "{{DATUM_VZNIKU}}": company_info.get("datum_vzniku", ""),
        "{{DATOVA_SCHRANKA}}": company_info.get("datova_schranka", ""),
        "{{DATUM_DNES}}": now.strftime("%d.%m.%Y"),
        "{{ROK_DNES}}": str(now.year),
        "{{SIDLO_KOMPLET}}": _build_full_address(company_info),
        "{{ZAKLADNI_KAPITAL}}": company_info.get("zakladni_kapital", ""),
        "{{PREDMET_PODNIKANI}}": company_info.get("predmet_podnikani", ""),
    }


def _build_full_address(info: dict) -> str:
    parts = []
    if info.get("sidlo_ulice"):
        parts.append(info["sidlo_ulice"])
    if info.get("sidlo_mesto"):
        parts.append(info["sidlo_mesto"])
    if info.get("sidlo_psc"):
        parts.append(info["sidlo_psc"])
    return ", ".join(parts)


def fill_template(template_path: str, placeholders: dict) -> bytes:
    """
    Načte Word šablonu, nahradí placeholdery a vrátí výsledný .docx jako bytes.
    """
    doc = Document(template_path)

    # Nahrazení v odstavcích
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, placeholders)

    # Nahrazení v tabulkách
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, placeholders)

    # Nahrazení v záhlaví a zápatí
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            if header_footer:
                for paragraph in header_footer.paragraphs:
                    _replace_in_paragraph(paragraph, placeholders)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _replace_in_paragraph(paragraph, placeholders: dict):
    """Nahradí placeholdery v odstavci, zachovává formátování."""
    full_text = paragraph.text
    if not any(key in full_text for key in placeholders):
        return

    # Jednoduchý přístup: nahradit v celém textu
    for key, value in placeholders.items():
        if key in full_text:
            full_text = full_text.replace(key, str(value))

    # Přepsat text – zachová formátování prvního runu
    if paragraph.runs:
        # Uložíme formátování prvního runu
        first_run = paragraph.runs[0]
        # Smažeme všechny runy kromě prvního
        for run in paragraph.runs[1:]:
            run.text = ""
        first_run.text = full_text
    else:
        paragraph.text = full_text


def generate_all_documents(company_info: dict, template_names: Optional[list[str]] = None) -> bytes:
    """
    Vygeneruje všechny (nebo vybrané) dokumenty a vrátí jako ZIP archiv.
    """
    placeholders = build_placeholders(company_info)
    templates = get_available_templates()

    if template_names:
        templates = [t for t in templates if t["name"] in template_names]

    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for tmpl in templates:
            try:
                doc_bytes = fill_template(tmpl["path"], placeholders)
                output_name = f"{tmpl['name']}_{company_info.get('ico', 'export')}.docx"
                zf.writestr(output_name, doc_bytes)
            except Exception:
                continue

    return buf.getvalue()


def create_sample_templates():
    """Vytvoří ukázkové .docx šablony s placeholdery."""
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

    templates = {
        "Smlouva_danove_poradenstvi": _create_smlouva_template,
        "Plna_moc": _create_plna_moc_template,
        "GDPR_souhlas": _create_gdpr_template,
        "OPDP_oznameni": _create_opdp_template,
    }

    for name, create_fn in templates.items():
        path = TEMPLATES_DIR / f"{name}.docx"
        if not path.exists():
            create_fn(path)


def _create_smlouva_template(path: Path):
    doc = Document()
    doc.add_heading("SMLOUVA O POSKYTOVÁNÍ DAŇOVÉHO PORADENSTVÍ", level=1)
    doc.add_paragraph("")
    doc.add_paragraph("uzavřená dle zákona č. 523/1992 Sb., o daňovém poradenství")
    doc.add_paragraph("")
    doc.add_heading("Smluvní strany", level=2)
    doc.add_paragraph("")
    doc.add_paragraph("1. Daňový poradce:")
    doc.add_paragraph("MDG tax & accounting s.r.o.")
    doc.add_paragraph("")
    doc.add_paragraph("2. Klient:")
    doc.add_paragraph("Název: {{NAZEV_FIRMY}}")
    doc.add_paragraph("IČO: {{ICO}}")
    doc.add_paragraph("DIČ: {{DIC}}")
    doc.add_paragraph("Sídlo: {{SIDLO_ULICE}}, {{SIDLO_MESTO}}, {{SIDLO_PSC}}")
    doc.add_paragraph("Právní forma: {{PRAVNI_FORMA}}")
    doc.add_paragraph("Zastoupen: {{JEDNATEL_JMENO}}, {{JEDNATEL_FUNKCE}}")
    doc.add_paragraph("Datová schránka: {{DATOVA_SCHRANKA}}")
    doc.add_paragraph("")
    doc.add_heading("Předmět smlouvy", level=2)
    doc.add_paragraph(
        "Daňový poradce se zavazuje poskytovat klientovi daňové poradenství "
        "v rozsahu stanoveném touto smlouvou, a to za podmínek a za odměnu "
        "sjednanou v této smlouvě."
    )
    doc.add_paragraph("")
    doc.add_heading("Závěrečná ustanovení", level=2)
    doc.add_paragraph("Tato smlouva nabývá platnosti a účinnosti dnem podpisu obou smluvních stran.")
    doc.add_paragraph(f"V Praze dne {{{{DATUM_DNES}}}}")
    doc.add_paragraph("")
    doc.add_paragraph("_________________________          _________________________")
    doc.add_paragraph("Daňový poradce                      Klient")
    doc.save(str(path))


def _create_plna_moc_template(path: Path):
    doc = Document()
    doc.add_heading("PLNÁ MOC", level=1)
    doc.add_paragraph("")
    doc.add_paragraph("Zmocnitel:")
    doc.add_paragraph("{{NAZEV_FIRMY}}")
    doc.add_paragraph("IČO: {{ICO}}, DIČ: {{DIC}}")
    doc.add_paragraph("se sídlem {{SIDLO_ULICE}}, {{SIDLO_MESTO}}, {{SIDLO_PSC}}")
    doc.add_paragraph("zastoupen {{JEDNATEL_JMENO}}, {{JEDNATEL_FUNKCE}}")
    doc.add_paragraph("")
    doc.add_paragraph("tímto zmocňuje")
    doc.add_paragraph("")
    doc.add_paragraph("Zmocněnce:")
    doc.add_paragraph("MDG tax & accounting s.r.o.")
    doc.add_paragraph("")
    doc.add_paragraph(
        "k zastupování zmocnitele před správcem daně ve všech daňových řízeních "
        "a při všech úkonech, ke kterým je zmocněnec oprávněn na základě "
        "zákona č. 523/1992 Sb., o daňovém poradenství."
    )
    doc.add_paragraph("")
    doc.add_paragraph(f"V Praze dne {{{{DATUM_DNES}}}}")
    doc.add_paragraph("")
    doc.add_paragraph("_________________________")
    doc.add_paragraph("Zmocnitel")
    doc.save(str(path))


def _create_gdpr_template(path: Path):
    doc = Document()
    doc.add_heading("INFORMACE O ZPRACOVÁNÍ OSOBNÍCH ÚDAJŮ", level=1)
    doc.add_paragraph("dle Nařízení Evropského parlamentu a Rady (EU) 2016/679 (GDPR)")
    doc.add_paragraph("")
    doc.add_paragraph("Správce osobních údajů:")
    doc.add_paragraph("MDG tax & accounting s.r.o.")
    doc.add_paragraph("")
    doc.add_paragraph("Subjekt údajů / Klient:")
    doc.add_paragraph("{{NAZEV_FIRMY}}, IČO: {{ICO}}")
    doc.add_paragraph("Sídlo: {{SIDLO_ULICE}}, {{SIDLO_MESTO}}, {{SIDLO_PSC}}")
    doc.add_paragraph("Zastoupen: {{JEDNATEL_JMENO}}")
    doc.add_paragraph("")
    doc.add_heading("Účel zpracování", level=2)
    doc.add_paragraph(
        "Osobní údaje jsou zpracovávány za účelem poskytování daňového poradenství, "
        "účetních služeb a plnění zákonných povinností dle zákona č. 253/2008 Sb. (AML zákon)."
    )
    doc.add_paragraph("")
    doc.add_heading("Doba zpracování", level=2)
    doc.add_paragraph(
        "Osobní údaje budou zpracovávány po dobu trvání smluvního vztahu a dále "
        "po dobu stanovenou příslušnými právními předpisy."
    )
    doc.add_paragraph("")
    doc.add_paragraph(f"Datum: {{{{DATUM_DNES}}}}")
    doc.save(str(path))


def _create_opdp_template(path: Path):
    doc = Document()
    doc.add_heading("OZNÁMENÍ O POVINNÉ DATOVÉ POVINNOSTI", level=1)
    doc.add_paragraph("")
    doc.add_paragraph("Klient:")
    doc.add_paragraph("{{NAZEV_FIRMY}}")
    doc.add_paragraph("IČO: {{ICO}}, DIČ: {{DIC}}")
    doc.add_paragraph("Sídlo: {{SIDLO_ULICE}}, {{SIDLO_MESTO}}, {{SIDLO_PSC}}")
    doc.add_paragraph("Datová schránka: {{DATOVA_SCHRANKA}}")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Tímto Vás informujeme o povinnosti činit podání vůči orgánům Finanční správy "
        "datovou zprávou ve smyslu § 72 odst. 4 zákona č. 280/2009 Sb., daňový řád."
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "Tato povinnost se vztahuje na všechna podání, u kterých tak stanoví zákon, "
        "zejména na daňová přiznání, hlášení a vyúčtování."
    )
    doc.add_paragraph("")
    doc.add_paragraph(f"V Praze dne {{{{DATUM_DNES}}}}")
    doc.add_paragraph("")
    doc.add_paragraph("MDG tax & accounting s.r.o.")
    doc.save(str(path))
